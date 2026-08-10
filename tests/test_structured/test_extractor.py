"""Integration tests for StructuredExtractor across PDF/XLSX/CSV/DOCX."""

import csv
import os
import tempfile

import pytest

from goblintools import TextExtractor
from goblintools.structured import StructuredExtractor


def _make_pdf_with_item_table(path: str) -> None:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle

    doc = SimpleDocTemplate(path, pagesize=letter)
    data = [
        ["ITEM", "DESCRICAO", "QTD", "UND"],
        ["1", "Notebook i7", "10", "UN"],
        ["2", "Mouse", "20", "UN"],
    ]
    table = Table(data, colWidths=[50, 200, 50, 50])
    table.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
                ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
            ]
        )
    )
    doc.build([table])


def _make_xlsx(path: str) -> None:
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Itens"
    ws.append(["ITEM", "DESCRICAO", "QTD", "UND"])
    ws.append(["1", "Caderno", "100", "UN"])
    ws.append(["2", "Caneta", "50", "UN"])
    wb.save(path)


def _make_csv(path: str) -> None:
    with open(path, "w", newline="", encoding="utf-8") as fh:
        writer = csv.writer(fh)
        writer.writerow(["ITEM", "DESCRICAO", "QTD", "UND"])
        writer.writerow(["1", "Papel A4", "10", "RESMA"])
        writer.writerow(["2", "Grampeador", "5", "UN"])


def _make_docx(path: str) -> None:
    import docx

    document = docx.Document()
    document.add_paragraph("Relacao de itens do edital")
    table = document.add_table(rows=3, cols=4)
    headers = ["ITEM", "DESCRICAO", "QTD", "UND"]
    data = [
        ["1", "Mesa", "4", "UN"],
        ["2", "Cadeira", "8", "UN"],
    ]
    for j, h in enumerate(headers):
        table.rows[0].cells[j].text = h
    for i, row in enumerate(data):
        for j, val in enumerate(row):
            table.rows[i + 1].cells[j].text = val
    document.save(path)


def _make_plain_pdf(path: str) -> None:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    c = canvas.Canvas(path, pagesize=letter)
    c.drawString(72, 720, "Hello plain PDF without tables")
    c.save()


@pytest.fixture
def item_pdf():
    pytest.importorskip("reportlab")
    pytest.importorskip("pdfplumber")
    with tempfile.TemporaryDirectory() as tmp:
        path = os.path.join(tmp, "itens.pdf")
        _make_pdf_with_item_table(path)
        yield path


@pytest.fixture
def plain_pdf():
    pytest.importorskip("reportlab")
    with tempfile.TemporaryDirectory() as tmp:
        path = os.path.join(tmp, "plain.pdf")
        _make_plain_pdf(path)
        yield path


def test_extractor_pdf_ok_for_items(item_pdf):
    doc = StructuredExtractor().extract_from_file(item_pdf)
    assert doc.tables, "expected at least one table from fixture PDF"
    assert doc.ok_for_items
    md = StructuredExtractor().to_full_md(doc)
    assert "<table>" in md
    assert "ITEM" in md


def test_extractor_xlsx_ok_for_items():
    with tempfile.TemporaryDirectory() as tmp:
        path = os.path.join(tmp, "itens.xlsx")
        _make_xlsx(path)
        doc = StructuredExtractor().extract_from_file(path)
        assert doc.ok_for_items
        assert doc.tables[0].sheet == "Itens"
        assert doc.tables[0].source == "xlsx"


def test_extractor_csv_ok_for_items():
    with tempfile.TemporaryDirectory() as tmp:
        path = os.path.join(tmp, "itens.csv")
        _make_csv(path)
        doc = StructuredExtractor().extract_from_file(path)
        assert doc.ok_for_items
        assert doc.tables[0].source == "csv"


def test_extractor_docx_ok_for_items():
    with tempfile.TemporaryDirectory() as tmp:
        path = os.path.join(tmp, "itens.docx")
        _make_docx(path)
        doc = StructuredExtractor().extract_from_file(path)
        assert doc.ok_for_items
        assert "Relacao de itens" in doc.prose
        md = StructuredExtractor().to_full_md(doc)
        assert "<table>" in md


def test_extractor_unsupported_suffix():
    with tempfile.NamedTemporaryFile(suffix=".txt", delete=False) as fh:
        fh.write(b"hello")
        path = fh.name
    try:
        doc = StructuredExtractor().extract_from_file(path)
        assert doc.tables == []
        assert not doc.ok_for_items
    finally:
        os.unlink(path)


def test_write_full_md(item_pdf):
    ext = StructuredExtractor()
    doc = ext.extract_from_file(item_pdf)
    if not doc.ok_for_items:
        pytest.skip("pdfplumber found no item table in fixture")
    with tempfile.TemporaryDirectory() as tmp:
        dest = os.path.join(tmp, "full.md")
        assert ext.write_full_md(doc, dest)
        content = open(dest, encoding="utf-8").read()
        assert "<table>" in content


def test_text_extractor_plain_unaffected(plain_pdf):
    """Regression: TextExtractor defaults must ignore structured module."""
    off = TextExtractor(extract_tables=False).extract_from_file(plain_pdf)
    default = TextExtractor().extract_from_file(plain_pdf)
    assert off == default
    assert "<!-- table" not in off
    assert "<table>" not in off


def test_extract_from_folder_mixed():
    with tempfile.TemporaryDirectory() as tmp:
        _make_csv(os.path.join(tmp, "a.csv"))
        _make_xlsx(os.path.join(tmp, "b.xlsx"))
        with open(os.path.join(tmp, "ignore.txt"), "w", encoding="utf-8") as fh:
            fh.write("skip me")
        docs = StructuredExtractor().extract_from_folder(tmp)
        assert len(docs) == 2
        assert all(d.ok_for_items for d in docs)
