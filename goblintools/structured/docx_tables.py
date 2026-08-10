"""DOCX structured table extraction via python-docx."""

from __future__ import annotations

import logging
from typing import List, Optional, Tuple

from goblintools.table_extractor import normalize_table_rows
from goblintools.structured.models import TableBlock
from goblintools.structured.quality import prepare_rows_for_output, score_table

logger = logging.getLogger(__name__)


def extract_docx_content(path: str) -> Tuple[List[TableBlock], str]:
    """Return ``(tables, prose)`` from a DOCX file."""
    try:
        import docx
    except ImportError as e:
        raise ImportError(
            "python-docx is required for DOCX structured extraction"
        ) from e

    try:
        document = docx.Document(path)
    except Exception as e:
        logger.error("Failed to open DOCX %s: %s", path, e)
        return [], ""

    prose_parts: List[str] = []
    for para in document.paragraphs:
        text = (para.text or "").strip()
        if text:
            prose_parts.append(text)
    prose = "\n".join(prose_parts)

    blocks: List[TableBlock] = []
    for index, table in enumerate(document.tables):
        raw: List[List[Optional[str]]] = []
        for row in table.rows:
            cells: List[Optional[str]] = []
            for cell in row.cells:
                value = (cell.text or "").strip()
                cells.append(value if value else None)
            if any(c for c in cells):
                raw.append(cells)
        rows = normalize_table_rows(raw)
        rows = prepare_rows_for_output(rows)
        if not rows:
            continue
        quality = score_table(rows)
        if not quality.meaningful:
            continue
        blocks.append(
            TableBlock(
                page=None,
                index=index,
                source="docx",
                sheet=None,
                rows=rows,
                quality=quality,
            )
        )

    return blocks, prose
